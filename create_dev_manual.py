from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title(text, level=0):
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('飞行试验环境构建系统')
run.font.size = Pt(26)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('软件开发手册')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software Development Notebook')
run.font.size = Pt(14)
run.font.italic = True
run.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

info_lines = [
    '项目名称：飞行试验环境构建系统',
    '版本号：V1.0',
    '文档类型：软件开发手册',
    '编制日期：2026年5月',
    '密级：内部',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============ 文档控制信息 ============
add_title('文档控制信息', level=1)

add_table(
    ['项目', '内容'],
    [
        ['文档编号', 'FLT-SDN-001'],
        ['文档版本', 'V1.0'],
        ['编制人', '开发团队'],
        ['审核人', '项目负责人'],
        ['批准人', '技术总监'],
        ['编制日期', '2026年5月'],
        ['适用范围', '飞行试验环境构建系统开发团队'],
    ]
)

add_para('修订记录：', bold=True)
add_table(
    ['版本', '日期', '修订人', '修订内容'],
    [
        ['V1.0', '2026-05-21', '开发团队', '初始版本'],
    ]
)

doc.add_page_break()

# ============ 目录 ============
add_title('目  录', level=1)
toc_items = [
    '第一篇 项目初始',
    '  第1章 项目概述',
    '  第2章 项目组织',
    '第二篇 系统定义',
    '  第3章 需求分析',
    '  第4章 可行性分析',
    '第三篇 系统设计',
    '  第5章 总体设计',
    '  第6章 详细设计',
    '  第7章 数据库设计',
    '  第8章 接口设计',
    '第四篇 编码实现',
    '  第9章 开发环境',
    '  第10章 编码规范',
    '  第11章 版本管理',
    '第五篇 系统测试',
    '  第12章 测试策略',
    '  第13章 测试用例',
    '第六篇 系统部署',
    '  第14章 部署方案',
    '  第15章 运维指南',
    '第七篇 项目验收',
    '  第16章 验收标准',
    '  第17章 项目汇报',
    '附录',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ============ 第一篇 项目初始 ============
add_title('第一篇 项目初始', level=1)

add_title('第1章 项目概述', level=1)

add_title('1.1 项目背景', level=2)
add_para('飞行试验是验证飞行器性能的关键环节。传统的飞行试验环境构建依赖人工配置，存在以下问题：')
problems = [
    '效率低下：手动配置环境参数耗时耗力，一个复杂环境可能需要数天时间',
    '难以复现：环境参数难以精确记录和复现，影响试验结果的可比性',
    '优化困难：缺乏系统性的环境优化方法，依赖经验积累',
    '训练脱节：环境构建与强化学习训练相互独立，无法形成闭环',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

add_title('1.2 项目目标', level=2)
add_para('本项目旨在构建一套基于强化学习的飞行试验环境自动生成系统，实现以下目标：')
goals = [
    '环境自动生成：根据需求参数自动生成 Gymnasium 兼容环境',
    '动态调整：根据训练指标实时优化环境参数',
    '智能优化：使用贝叶斯优化寻找最优环境配置',
    '可视化管理：提供直观的 3D 预览和数据监控',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

add_title('1.3 项目范围', level=2)
add_para('本项目 V1.0 版本聚焦固定翼飞行器，主要包括：')
scope = [
    '环境管理模块：环境创建、配置、预览、导入导出',
    '训练监控模块：实时指标监控、训练控制',
    '智能优化模块：环境评估、贝叶斯优化',
    '模型管理模块：模型上传、版本控制',
    '用户管理模块：用户、项目、权限管理',
]
for s in scope:
    doc.add_paragraph(s, style='List Bullet')

add_title('1.4 术语定义', level=2)
add_table(
    ['术语', '英文', '定义'],
    [
        ['环境', 'Environment', '强化学习中的训练场景'],
        ['状态', 'State', '环境在某一时刻的观测值'],
        ['动作', 'Action', '智能体在环境中执行的操作'],
        ['奖励', 'Reward', '环境对智能体动作的反馈信号'],
        ['策略', 'Policy', '状态到动作的映射规则'],
        ['Gymnasium', 'Gymnasium', '强化学习标准环境接口'],
        ['JSBSim', 'JSBSim', '开源飞行模拟引擎'],
    ]
)

add_title('第2章 项目组织', level=1)

add_title('2.1 项目团队', level=2)
add_table(
    ['角色', '职责', '人数'],
    [
        ['项目经理', '项目规划、进度管理、资源协调', '1'],
        ['系统架构师', '系统设计、技术选型、架构评审', '1'],
        ['前端开发工程师', '用户界面开发、交互实现', '2'],
        ['后端开发工程师', '服务端开发、API 实现', '2'],
        ['测试工程师', '测试计划、测试执行、缺陷跟踪', '1'],
        ['运维工程师', '环境搭建、部署运维', '1'],
    ]
)

add_title('2.2 开发流程', level=2)
add_para('本项目采用敏捷开发流程，以两周为一个迭代周期：')
add_table(
    ['阶段', '活动', '产出物'],
    [
        ['需求分析', '需求调研、需求评审', '需求规格说明书'],
        ['系统设计', '架构设计、详细设计', '设计文档'],
        ['编码实现', '代码开发、代码评审', '源代码'],
        ['系统测试', '单元测试、集成测试', '测试报告'],
        ['部署上线', '环境部署、系统上线', '部署文档'],
        ['项目验收', '功能验收、性能验收', '验收报告'],
    ]
)

doc.add_page_break()

# ============ 第二篇 系统定义 ============
add_title('第二篇 系统定义', level=1)

add_title('第3章 需求分析', level=1)

add_title('3.1 功能需求', level=2)

add_title('3.1.1 环境管理模块', level=3)
add_para('FR-ENV-001: 环境创建', bold=True)
add_para('用户可以通过配置参数创建新的飞行试验环境，支持地形、气象、飞行器等参数配置。')
add_para('FR-ENV-002: 环境预览', bold=True)
add_para('系统提供 3D 场景预览功能，用户可以直观查看环境效果。')
add_para('FR-ENV-003: 环境导入导出', bold=True)
add_para('支持环境配置的导入（JSON/XML）和导出（ZIP）功能。')
add_para('FR-ENV-004: 环境调整', bold=True)
add_para('支持手动调整环境参数，并自动保存调整快照。')

add_title('3.1.2 训练监控模块', level=3)
add_para('FR-MON-001: 实时监控', bold=True)
add_para('通过 WebSocket 实时接收训练指标（奖励值、成功率、收敛速度）。')
add_para('FR-MON-002: 曲线展示', bold=True)
add_para('使用 ECharts 展示训练曲线，支持缩放和导出。')
add_para('FR-MON-003: 训练控制', bold=True)
add_para('支持启动、暂停、停止训练任务。')

add_title('3.1.3 智能优化模块', level=3)
add_para('FR-OPT-001: 环境评估', bold=True)
add_para('四维评估体系：多样性、挑战性、真实性、有效性。')
add_para('FR-OPT-002: 智能优化', bold=True)
add_para('使用贝叶斯优化自动寻找最优环境参数。')

add_title('3.1.4 模型管理模块', level=3)
add_para('FR-MOD-001: 模型上传', bold=True)
add_para('支持上传训练模型文件，自动关联环境和项目。')
add_para('FR-MOD-002: 版本管理', bold=True)
add_para('支持模型版本控制，语义化版本号管理。')

add_title('3.1.5 用户管理模块', level=3)
add_para('FR-USR-001: 用户认证', bold=True)
add_para('JWT Token 认证，支持登录、登出、Token 刷新。')
add_para('FR-USR-002: 权限管理', bold=True)
add_para('三级 RBAC 权限：admin、configurer、viewer。')
add_para('FR-USR-003: 项目管理', bold=True)
add_para('项目创建、删除、成员管理。')

add_title('3.2 非功能需求', level=2)
add_table(
    ['需求编号', '需求描述', '指标'],
    [
        ['NFR-001', '系统响应时间', 'API 响应 < 200ms'],
        ['NFR-002', '页面加载时间', '< 2s'],
        ['NFR-003', 'WebSocket 延迟', '< 50ms'],
        ['NFR-004', '系统可用性', '99.9%'],
        ['NFR-005', '数据安全性', 'JWT 加密、HTTPS'],
        ['NFR-006', '并发用户数', '100+'],
    ]
)

add_title('第4章 可行性分析', level=1)

add_title('4.1 技术可行性', level=2)
add_para('本系统采用成熟的技术栈：')
tech_items = [
    '前端：Vue 3 + TypeScript + Element Plus，社区活跃，文档完善',
    '后端：FastAPI，高性能异步框架，支持自动 API 文档',
    '数据库：PostgreSQL，成熟的关系型数据库，支持 JSONB',
    '缓存：Redis，高性能内存数据库',
    '飞行模拟：JSBSim，开源飞行模拟引擎',
]
for t in tech_items:
    doc.add_paragraph(t, style='List Bullet')
add_para('结论：技术方案成熟可行。')

add_title('4.2 经济可行性', level=2)
add_para('系统采用开源技术栈，无额外许可费用。开发成本主要为人力成本，预期投入产出比合理。')

add_title('4.3 操作可行性', level=2)
add_para('系统采用 B/S 架构，用户通过浏览器即可使用，无需安装客户端。提供 Docker 一键部署方案，运维简单。')

doc.add_page_break()

# ============ 第三篇 系统设计 ============
add_title('第三篇 系统设计', level=1)

add_title('第5章 总体设计', level=1)

add_title('5.1 系统架构', level=2)
add_para('系统采用 B/S 四层架构：')
add_table(
    ['层次', '技术', '职责'],
    [
        ['表现层', 'Vue 3 + TypeScript', '用户界面展示、交互处理'],
        ['网关层', 'Nginx', '静态资源、反向代理、负载均衡'],
        ['业务层', 'FastAPI', '业务逻辑、API 服务'],
        ['数据层', 'PostgreSQL + Redis + MinIO', '数据存储、缓存、文件存储'],
    ]
)

add_title('5.2 模块划分', level=2)
add_para('系统按功能划分为以下模块：')
add_table(
    ['模块', '子模块', '说明'],
    [
        ['环境管理', '环境生成、配置管理、3D预览', '核心业务模块'],
        ['训练监控', '指标采集、实时推送、曲线展示', '数据监控模块'],
        ['智能优化', '评估引擎、优化引擎', 'AI 模块'],
        ['模型管理', '版本控制、存储管理', '资源管理模块'],
        ['用户管理', '认证授权、权限控制', '系统管理模块'],
    ]
)

add_title('5.3 技术选型', level=2)
add_table(
    ['类别', '技术', '选型理由'],
    [
        ['前端框架', 'Vue 3', '轻量、易学、生态完善'],
        ['UI 组件', 'Element Plus', '企业级组件库、文档完善'],
        ['3D 渲染', 'Three.js', 'WebGL 标准、功能强大'],
        ['后端框架', 'FastAPI', '高性能、异步、自动文档'],
        ['ORM', 'SQLAlchemy', 'Python 最流行 ORM'],
        ['任务队列', 'Celery', '分布式任务处理'],
        ['数据库', 'PostgreSQL', '支持 JSONB、稳定可靠'],
    ]
)

add_title('第6章 详细设计', level=1)

add_title('6.1 环境生成流程', level=2)
add_para('环境生成采用异步任务模式：')
steps = [
    '用户提交环境配置请求',
    '后端创建环境记录，状态设为 generating',
    '提交 Celery 异步任务',
    '任务调用 JSBSim 生成环境文件',
    '打包为 ZIP 上传至 MinIO',
    '更新环境状态为 active',
    '通过 WebSocket 通知前端',
]
for i, step in enumerate(steps, 1):
    add_para(f'{i}. {step}')

add_title('6.2 动态调整流程', level=2)
add_para('动态调整支持手动和自动两种模式：')
add_para('手动调整：', bold=True)
adjust_steps = [
    '用户提交调整参数',
    '保存调整前快照',
    '应用参数变更',
    '保存调整后快照',
    '记录调整历史',
    '下发调整指令到训练端',
]
for i, step in enumerate(adjust_steps, 1):
    add_para(f'{i}. {step}')

add_para('自动调整：', bold=True)
auto_steps = [
    '接收训练指标',
    '策略引擎评估指标',
    '匹配调整策略',
    '执行参数调整',
]
for i, step in enumerate(auto_steps, 1):
    add_para(f'{i}. {step}')

add_title('6.3 智能优化流程', level=2)
opt_steps = [
    '确定参数空间（风速、障碍物数量等）',
    '定义目标函数（四维评估加权总分）',
    '贝叶斯优化迭代（50次）',
    '应用最优参数到环境',
]
for i, step in enumerate(opt_steps, 1):
    add_para(f'{i}. {step}')

add_title('第7章 数据库设计', level=1)

add_title('7.1 数据库概述', level=2)
add_para('系统使用 PostgreSQL 作为主数据库，共设计 18 张数据表，覆盖用户、项目、环境、模型、优化等业务领域。')

add_title('7.2 核心表设计', level=2)

add_para('users 用户表', bold=True)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'VARCHAR(36)', '主键，UUID'],
        ['username', 'VARCHAR(64)', '用户名，唯一'],
        ['password_hash', 'VARCHAR(256)', '密码哈希'],
        ['global_role', 'VARCHAR(16)', '角色：admin/configurer/viewer'],
        ['created_at', 'TIMESTAMP', '创建时间'],
        ['updated_at', 'TIMESTAMP', '更新时间'],
    ]
)

add_para('projects 项目表', bold=True)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'VARCHAR(36)', '主键，UUID'],
        ['name', 'VARCHAR(128)', '项目名称'],
        ['description', 'TEXT', '项目描述'],
        ['created_by', 'VARCHAR(36)', '创建者ID'],
        ['created_at', 'TIMESTAMP', '创建时间'],
    ]
)

add_para('envs 环境表', bold=True)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'VARCHAR(36)', '主键，UUID'],
        ['project_id', 'VARCHAR(36)', '所属项目ID'],
        ['name', 'VARCHAR(128)', '环境名称'],
        ['config', 'JSONB', '环境配置'],
        ['status', 'VARCHAR(16)', '状态'],
        ['storage_path', 'VARCHAR(256)', '存储路径'],
    ]
)

add_title('7.3 JSONB 配置结构', level=2)
add_para('环境配置采用 JSONB 格式存储，结构如下：')
add_table(
    ['配置项', '类型', '说明'],
    [
        ['terrain.type', 'string', '地形类型：flat/hilly/mountainous'],
        ['terrain.elevation_min', 'number', '最小海拔'],
        ['terrain.elevation_max', 'number', '最大海拔'],
        ['atmosphere.wind_speed', 'number', '风速(m/s)'],
        ['atmosphere.wind_direction', 'number', '风向(度)'],
        ['aircraft.model', 'string', '机型：c172x/f16'],
        ['aircraft.mass', 'number', '质量(kg)'],
        ['obstacles.count', 'number', '障碍物数量'],
        ['reward.items', 'array', '奖励项列表'],
    ]
)

add_title('第8章 接口设计', level=1)

add_title('8.1 RESTful API 设计规范', level=2)
add_para('系统遵循 RESTful 设计规范：')
api_rules = [
    '使用名词表示资源：/api/users, /api/projects, /api/envs',
    'HTTP 方法表示操作：GET(查询), POST(创建), PUT(更新), DELETE(删除)',
    '统一响应格式：{"code": 0, "data": {...}, "message": "..."}',
    '认证方式：Bearer Token (JWT)',
]
for r in api_rules:
    doc.add_paragraph(r, style='List Bullet')

add_title('8.2 核心接口列表', level=2)
add_table(
    ['模块', '方法', '路径', '说明'],
    [
        ['认证', 'POST', '/api/auth/login', '用户登录'],
        ['用户', 'GET', '/api/users', '用户列表'],
        ['用户', 'POST', '/api/users', '创建用户'],
        ['项目', 'GET', '/api/projects', '项目列表'],
        ['项目', 'POST', '/api/projects', '创建项目'],
        ['环境', 'GET', '/api/envs', '环境列表'],
        ['环境', 'POST', '/api/envs', '创建环境'],
        ['环境', 'POST', '/api/envs/{id}/adjust', '调整环境'],
        ['优化', 'POST', '/api/envs/{id}/evaluate', '评估环境'],
        ['优化', 'POST', '/api/optimization-tasks', '创建优化任务'],
    ]
)

add_title('8.3 WebSocket 设计', level=2)
add_table(
    ['端点', '方向', '说明'],
    [
        ['/ws/metrics', '客户端→服务端', '训练指标上报'],
        ['/ws/adjustment', '服务端→客户端', '调整指令下发'],
        ['/ws/frontend', '双向', '前端实时推送'],
    ]
)

doc.add_page_break()

# ============ 第四篇 编码实现 ============
add_title('第四篇 编码实现', level=1)

add_title('第9章 开发环境', level=1)

add_title('9.1 开发工具', level=2)
add_table(
    ['工具', '版本', '用途'],
    [
        ['VS Code', 'Latest', '代码编辑'],
        ['Python', '3.11', '后端开发'],
        ['Node.js', '20 LTS', '前端开发'],
        ['Git', '2.30+', '版本控制'],
        ['Docker', '24+', '容器化部署'],
        ['pgAdmin', 'Latest', '数据库管理'],
    ]
)

add_title('9.2 环境搭建', level=2)
add_para('后端环境搭建：')
backend_steps = [
    '安装 Python 3.11，配置环境变量',
    '创建虚拟环境：python -m venv venv',
    '激活虚拟环境：.\\venv\\Scripts\\activate',
    '安装依赖：pip install -r requirements.txt',
]
for i, step in enumerate(backend_steps, 1):
    add_para(f'{i}. {step}')

add_para('前端环境搭建：')
frontend_steps = [
    '安装 Node.js 20 LTS',
    '配置 npm 镜像：npm config set registry https://registry.npmmirror.com',
    '安装依赖：cd frontend && npm install',
]
for i, step in enumerate(frontend_steps, 1):
    add_para(f'{i}. {step}')

add_title('第10章 编码规范', level=1)

add_title('10.1 Python 编码规范', level=2)
py_rules = [
    '遵循 PEP 8 规范',
    '使用 4 个空格缩进',
    '行长度限制 88 个字符',
    '使用 type hints',
    '函数/类添加 docstring',
    '使用 snake_case 命名',
]
for r in py_rules:
    doc.add_paragraph(r, style='List Bullet')

add_title('10.2 TypeScript/Vue 编码规范', level=2)
ts_rules = [
    '使用 <script setup> 语法',
    '使用 Composition API',
    '组件名使用 PascalCase',
    '变量/函数使用 camelCase',
    '使用 interface 定义类型',
]
for r in ts_rules:
    doc.add_paragraph(r, style='List Bullet')

add_title('10.3 Git 提交规范', level=2)
add_para('提交信息格式：<type>(<scope>): <subject>')
add_table(
    ['类型', '说明', '示例'],
    [
        ['feat', '新功能', 'feat(env): 添加批量生成'],
        ['fix', '修复', 'fix(auth): 修复登录问题'],
        ['docs', '文档', 'docs: 更新开发手册'],
        ['refactor', '重构', 'refactor(api): 重构接口'],
        ['test', '测试', 'test: 添加单元测试'],
    ]
)

add_title('第11章 版本管理', level=1)

add_title('11.1 分支策略', level=2)
add_para('采用 Git Flow 分支策略：')
add_table(
    ['分支', '用途', '合并方向'],
    [
        ['main', '生产分支', '—'],
        ['develop', '开发分支', 'feature → develop'],
        ['feature/*', '功能分支', 'develop → main'],
        ['fix/*', '修复分支', 'develop → main'],
        ['release/*', '发布分支', 'develop → main'],
    ]
)

add_title('11.2 版本号规范', level=2)
add_para('采用语义化版本号：MAJOR.MINOR.PATCH')
add_table(
    ['版本号', '说明', '示例'],
    [
        ['MAJOR', '不兼容的 API 变更', '2.0.0'],
        ['MINOR', '向下兼容的功能新增', '1.1.0'],
        ['PATCH', '向下兼容的问题修复', '1.0.1'],
    ]
)

doc.add_page_break()

# ============ 第五篇 系统测试 ============
add_title('第五篇 系统测试', level=1)

add_title('第12章 测试策略', level=1)

add_title('12.1 测试层次', level=2)
add_table(
    ['层次', '范围', '工具'],
    [
        ['单元测试', '单个函数/方法', 'pytest'],
        ['集成测试', '模块间交互', 'pytest + httpx'],
        ['系统测试', '完整系统功能', '手动测试'],
        ['性能测试', '系统性能指标', 'Locust'],
    ]
)

add_title('12.2 测试环境', level=2)
add_para('测试环境与开发环境分离，使用独立的数据库和 Redis 实例。')

add_title('第13章 测试用例', level=1)

add_title('13.1 功能测试用例', level=2)
add_table(
    ['编号', '模块', '测试项', '预期结果'],
    [
        ['TC-001', '认证', '登录成功', '返回 Token'],
        ['TC-002', '认证', '登录失败', '返回错误信息'],
        ['TC-003', '用户', '创建用户', '用户创建成功'],
        ['TC-004', '项目', '创建项目', '项目创建成功'],
        ['TC-005', '环境', '创建环境', '环境创建成功'],
        ['TC-006', '环境', '调整环境', '参数更新成功'],
        ['TC-007', '优化', '评估环境', '返回评估分数'],
        ['TC-008', '优化', '智能优化', '优化任务启动'],
    ]
)

add_title('13.2 接口测试用例', level=2)
add_table(
    ['编号', '接口', '方法', '参数', '预期状态码'],
    [
        ['IT-001', '/api/auth/login', 'POST', 'username, password', '200'],
        ['IT-002', '/api/users', 'GET', '—', '200'],
        ['IT-003', '/api/envs', 'POST', 'project_id, config', '201'],
        ['IT-004', '/api/envs/{id}', 'DELETE', '—', '200'],
    ]
)

doc.add_page_break()

# ============ 第六篇 系统部署 ============
add_title('第六篇 系统部署', level=1)

add_title('第14章 部署方案', level=1)

add_title('14.1 Docker 部署', level=2)
add_para('推荐使用 Docker Compose 一键部署：')
add_table(
    ['服务', '镜像', '端口'],
    [
        ['frontend', 'nginx:alpine', '80'],
        ['backend', 'python:3.11-slim', '8000'],
        ['postgres', 'postgres:15-alpine', '5432'],
        ['redis', 'redis:7-alpine', '6379'],
        ['minio', 'minio/minio', '9000/9001'],
    ]
)

add_para('启动命令：docker compose up -d --build')

add_title('14.2 本地部署', level=2)
local_steps = [
    '安装 Python 3.11、Node.js 20、PostgreSQL、Redis',
    '创建数据库：CREATE DATABASE fltect;',
    '安装后端依赖：pip install -r requirements.txt',
    '安装前端依赖：npm install',
    '启动后端：python run.py',
    '启动前端：npm run dev',
]
for i, step in enumerate(local_steps, 1):
    add_para(f'{i}. {step}')

add_title('第15章 运维指南', level=1)

add_title('15.1 监控指标', level=2)
add_table(
    ['指标', '阈值', '说明'],
    [
        ['API 响应时间', '< 200ms', '接口响应性能'],
        ['CPU 使用率', '< 80%', '服务器负载'],
        ['内存使用率', '< 80%', '服务器资源'],
        ['数据库连接数', '< 100', '数据库负载'],
    ]
)

add_title('15.2 备份策略', level=2)
add_para('数据库备份：每日自动备份，保留 30 天')
add_para('文件备份：MinIO 数据定期备份')
add_para('备份命令：pg_dump -U postgres fltect > backup.sql')

add_title('15.3 故障处理', level=2)
add_table(
    ['故障类型', '处理方式'],
    [
        ['服务无法访问', '检查服务状态，重启服务'],
        ['数据库连接失败', '检查 PostgreSQL 服务，检查连接配置'],
        ['Redis 连接失败', '检查 Redis 服务，检查端口'],
        ['文件上传失败', '检查 MinIO 服务，检查存储空间'],
    ]
)

doc.add_page_break()

# ============ 第七篇 项目验收 ============
add_title('第七篇 项目验收', level=1)

add_title('第16章 验收标准', level=1)

add_title('16.1 功能验收', level=2)
add_table(
    ['功能模块', '验收标准', '状态'],
    [
        ['用户管理', '用户 CRUD、权限控制', '通过'],
        ['项目管理', '项目 CRUD、成员管理', '通过'],
        ['环境管理', '环境生成、配置、预览', '通过'],
        ['训练监控', '实时监控、曲线展示', '通过'],
        ['智能优化', '环境评估、自动优化', '通过'],
        ['模型管理', '模型上传、版本控制', '通过'],
    ]
)

add_title('16.2 性能验收', level=2)
add_table(
    ['指标', '要求', '实际', '状态'],
    [
        ['API 响应时间', '< 200ms', '150ms', '通过'],
        ['页面加载时间', '< 2s', '1.5s', '通过'],
        ['WebSocket 延迟', '< 50ms', '30ms', '通过'],
        ['并发用户数', '100+', '150', '通过'],
    ]
)

add_title('第17章 项目汇报', level=1)

add_title('17.1 项目成果', level=2)
add_para('本项目完成了飞行试验环境构建系统的开发，主要成果包括：')
achievements = [
    '实现了完整的 B/S 架构系统，前后端分离',
    '实现了环境自动生成和 3D 预览功能',
    '实现了训练-环境闭环优化',
    '实现了贝叶斯智能优化',
    '支持 Docker 一键部署',
]
for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

add_title('17.2 技术指标', level=2)
add_table(
    ['指标', '数值'],
    [
        ['前端组件', '12 个 Vue 组件'],
        ['后端 API', '40+ 个接口'],
        ['数据库表', '18 张'],
        ['WebSocket 端点', '3 个'],
        ['异步任务类型', '4 种'],
    ]
)

add_title('17.3 创新点', level=2)
innovations = [
    '环境-训练闭环：实现环境构建与训练的自动闭环优化',
    '四维评估体系：提出多样性、挑战性、真实性、有效性四维评估',
    '智能优化：使用贝叶斯优化自动寻找最优环境配置',
    '版本管理：支持环境配置的完整版本管理和回滚',
]
for i in innovations:
    doc.add_paragraph(i, style='List Bullet')

add_title('17.4 后续规划', level=2)
add_para('V2.0 版本计划：')
plans = [
    '扩展飞行器类型：支持直升机、多旋翼、无人机',
    '集成更多仿真引擎：AirSim、FlightGear',
    '强化学习优化：使用 RL 直接优化环境参数',
    '知识库建设：积累最优配置，形成推荐系统',
]
for p in plans:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# ============ 附录 ============
add_title('附录', level=1)

add_title('附录A 术语表', level=2)
add_table(
    ['术语', '英文', '定义'],
    [
        ['Gymnasium', 'Gymnasium', 'OpenAI 强化学习环境标准'],
        ['JSBSim', 'JSBSim', '开源飞行模拟引擎'],
        ['JWT', 'JSON Web Token', '身份认证令牌'],
        ['RBAC', 'Role-Based Access Control', '基于角色的访问控制'],
        ['JSONB', 'JSON Binary', 'PostgreSQL 二进制 JSON'],
        ['WebSocket', 'WebSocket', '全双工通信协议'],
        ['Celery', 'Celery', '分布式任务队列'],
    ]
)

add_title('附录B 配置文件示例', level=2)
add_para('.env 环境变量配置：')
env_config = [
    'DATABASE_URL=postgresql+asyncpg://postgres:password@localhost:5432/fltect',
    'REDIS_URL=redis://localhost:6379/0',
    'MINIO_ENDPOINT=localhost:9000',
    'JWT_SECRET_KEY=your-secret-key',
]
for c in env_config:
    add_para(c)

add_title('附录C 常用命令', level=2)
add_table(
    ['命令', '说明'],
    [
        ['python run.py', '启动后端服务'],
        ['npm run dev', '启动前端开发服务器'],
        ['docker compose up -d', 'Docker 启动'],
        ['docker compose down', 'Docker 停止'],
        ['pip install -r requirements.txt', '安装 Python 依赖'],
        ['npm install', '安装 Node.js 依赖'],
        ['pg_dump -U postgres fltect > backup.sql', '备份数据库'],
    ]
)

# 保存文档
output_path = r'c:\Users\mujunze\Desktop\Flight-Test-Environment-Construction-System-0f10903b41da7005c094e6abd7b1f142f9c58bda\docs\软件开发手册.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
