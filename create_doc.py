from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(9)  # 小五
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title(text, level=0):
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('飞行试验环境构建系统')
run.font.size = Pt(26)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('开发手册')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(6):
    doc.add_paragraph()

info_lines = [
    '版本：V1.0',
    '日期：2026年5月',
    '文档类型：开发手册',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============ 目录页 ============
add_title('目录', level=1)
toc_items = [
    '第一章 概述',
    '  1.1 文档目的',
    '  1.2 项目简介',
    '  1.3 技术栈概览',
    '  1.4 系统架构',
    '第二章 开发环境搭建',
    '  2.1 操作系统要求',
    '  2.2 Python 环境配置',
    '  2.3 Node.js 环境配置',
    '  2.4 IDE 配置',
    '  2.5 数据库环境配置',
    '第三章 项目结构详解',
    '  3.1 整体目录结构',
    '  3.2 后端目录详解',
    '  3.3 前端目录详解',
    '第四章 开发规范',
    '  4.1 Python 代码规范',
    '  4.2 TypeScript/Vue 代码规范',
    '  4.3 Git 规范',
    '第五章 调试与测试',
    '  5.1 后端调试',
    '  5.2 前端调试',
    '  5.3 数据库调试',
    '第六章 部署指南',
    '  6.1 Docker 部署',
    '  6.2 本地部署',
    '第七章 常见问题',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ============ 第一章 概述 ============
add_title('第一章 概述', level=1)

add_title('1.1 文档目的', level=2)
add_para('本文档旨在为开发人员提供飞行试验环境构建系统的完整开发指南，包括环境搭建、配置说明、开发规范等内容。通过阅读本文档，开发人员可以快速上手项目开发，理解系统架构，并能够独立完成模块开发和调试工作。')

add_title('1.2 项目简介', level=2)
add_para('飞行试验环境构建系统是一套基于强化学习的飞行试验环境自动生成平台。系统根据用户需求自动生成 Gymnasium 兼容的飞行试验环境，支持动态调整和智能优化。V1.0 版本聚焦固定翼飞行器。')
add_para('系统主要功能包括：')
features = [
    '环境自动生成：根据配置参数自动构建 Gymnasium 兼容环境',
    '动态调整：根据训练指标实时优化环境参数',
    '智能优化：使用贝叶斯优化寻找最优环境配置',
    '可视化管理：提供直观的 3D 预览和数据监控',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_title('1.3 技术栈概览', level=2)
add_table(
    ['层级', '技术', '版本', '用途'],
    [
        ['前端框架', 'Vue', '3.5+', '响应式 UI 框架'],
        ['UI 组件库', 'Element Plus', '2.14+', '企业级 UI 组件'],
        ['3D 渲染', 'Three.js', '0.184+', '三维场景预览'],
        ['数据可视化', 'ECharts', '6.0+', '图表展示'],
        ['后端框架', 'FastAPI', '0.110+', '异步 Web 框架'],
        ['ORM', 'SQLAlchemy', '2.0+', '数据库操作'],
        ['异步任务', 'Celery', '5.3+', '分布式任务队列'],
        ['缓存/消息', 'Redis', '7.0+', '缓存和消息代理'],
        ['数据库', 'PostgreSQL', '15+', '关系型数据库'],
        ['对象存储', 'MinIO', 'Latest', '文件存储'],
        ['版本控制', 'Git', '2.30+', '代码版本管理'],
        ['容器化', 'Docker', '24+', '应用容器化'],
    ]
)

add_title('1.4 系统架构', level=2)
add_para('系统采用 B/S 四层架构设计，包括表现层、网关层、业务层和数据层。')
add_para('表现层：Vue 3 + TypeScript + Element Plus + Three.js + ECharts', bold=True)
add_para('负责用户界面展示，包括环境配置、3D 预览、训练监控、数据可视化等功能。')
add_para('网关层：Nginx', bold=True)
add_para('负责静态资源托管、API 反向代理、WebSocket 转发。')
add_para('业务层：FastAPI', bold=True)
add_para('核心业务逻辑，包括环境生成、动态调整、智能优化、模型管理等模块。')
add_para('数据层：PostgreSQL + Redis + MinIO', bold=True)
add_para('PostgreSQL 存储业务数据，Redis 用作缓存和消息队列，MinIO 存储文件。')

doc.add_page_break()

# ============ 第二章 开发环境搭建 ============
add_title('第二章 开发环境搭建', level=1)

add_title('2.1 操作系统要求', level=2)
add_table(
    ['操作系统', '版本要求', '备注'],
    [
        ['Windows', '10/11 64位', '推荐开发环境'],
        ['macOS', '12.0+', '需要额外配置'],
        ['Ubuntu', '20.04+', '推荐服务器部署'],
        ['CentOS', '8+', '推荐服务器部署'],
    ]
)

add_title('2.2 Python 环境配置', level=2)

add_title('2.2.1 安装 Python', level=3)
steps = [
    '访问 Python 官方网站：https://www.python.org/downloads/',
    '下载 Python 3.11 版本（推荐）',
    '运行安装程序，务必勾选 Add Python 3.11 to PATH',
    '安装完成后验证：python --version',
]
for i, step in enumerate(steps, 1):
    add_para(f'{i}. {step}')

add_title('2.2.2 配置 pip 镜像', level=3)
add_para('为了加速依赖包下载，配置国内镜像源：')
add_para('pip config set global.index-url https://pypi.tuna.tsinghua.edu.cn/simple')

add_title('2.2.3 创建虚拟环境', level=3)
add_para('虚拟环境用于隔离项目依赖：')
cmds = [
    'cd backend',
    'python -m venv venv',
    '.\\venv\\Scripts\\activate  # Windows',
    'source venv/bin/activate   # Linux/Mac',
]
for cmd in cmds:
    add_para(cmd)

add_title('2.2.4 安装后端依赖', level=3)
add_para('pip install -r requirements.txt')
add_para('核心依赖说明：')
add_table(
    ['包名', '版本要求', '用途说明'],
    [
        ['fastapi', '>=0.110', '高性能异步 Web 框架'],
        ['uvicorn', '>=0.29', 'ASGI 服务器'],
        ['sqlalchemy', '>=2.0', '异步 ORM 框架'],
        ['asyncpg', '>=0.29', 'PostgreSQL 异步驱动'],
        ['celery', '>=5.3', '分布式任务队列'],
        ['redis', '>=5.0', 'Redis 客户端'],
        ['python-jose', '>=3.3', 'JWT 令牌处理'],
        ['passlib', '>=1.7', '密码哈希加密'],
        ['jsbsim', '>=1.1', '飞行模拟引擎'],
        ['scikit-optimize', '>=0.9', '贝叶斯优化'],
    ]
)

add_title('2.3 Node.js 环境配置', level=2)

add_title('2.3.1 安装 Node.js', level=3)
add_para('1. 访问 Node.js 官方网站：https://nodejs.org/')
add_para('2. 下载 LTS 版本（推荐 20.x）')
add_para('3. 运行安装程序，使用默认配置安装')
add_para('4. 验证安装：node --version')

add_title('2.3.2 配置 npm 镜像', level=3)
add_para('npm config set registry https://registry.npmmirror.com')

add_title('2.3.3 安装前端依赖', level=3)
add_para('cd frontend && npm install')
add_para('前端依赖说明：')
add_table(
    ['包名', '版本', '用途说明'],
    [
        ['vue', '^3.5.34', 'Vue 3 核心框架'],
        ['vue-router', '^5.0.7', '路由管理'],
        ['pinia', '^3.0.4', '状态管理'],
        ['element-plus', '^2.14.0', 'UI 组件库'],
        ['axios', '^1.16.1', 'HTTP 请求库'],
        ['echarts', '^6.0.0', '数据可视化'],
        ['three', '^0.184.0', '3D 渲染引擎'],
    ]
)

add_title('2.4 IDE 配置', level=2)

add_title('2.4.1 VS Code 配置', level=3)
add_para('安装必要扩展：')
add_table(
    ['扩展名称', '发布者', '用途'],
    [
        ['Python', 'Microsoft', 'Python 语言支持'],
        ['Pylance', 'Microsoft', 'Python 智能提示'],
        ['Vue - Official', 'Vue', 'Vue 3 语言支持'],
        ['ESLint', 'Microsoft', '代码检查'],
        ['Prettier', 'Prettier', '代码格式化'],
        ['GitLens', 'GitLens', 'Git 增强'],
    ]
)

add_title('2.5 数据库环境配置', level=2)

add_title('2.5.1 PostgreSQL 安装', level=3)
add_para('1. 访问 https://www.postgresql.org/download/windows/')
add_para('2. 下载 PostgreSQL 15 安装程序')
add_para('3. 安装时设置密码（务必记住）')
add_para('4. 默认端口：5432')

add_title('2.5.2 创建数据库', level=3)
add_para('连接 PostgreSQL 后执行：')
add_para('CREATE DATABASE fltect;')

add_title('2.5.3 配置连接', level=3)
add_para('编辑 backend/app/core/config.py：')
add_para('DATABASE_URL = "postgresql+asyncpg://postgres:your_password@localhost:5432/fltect"')

add_title('2.5.4 Redis 配置', level=3)
add_para('1. 下载 Redis: https://github.com/tporadowski/redis/releases')
add_para('2. 解压后运行 redis-server.exe')
add_para('3. 验证：redis-cli ping → PONG')

add_title('2.5.5 MinIO 配置（可选）', level=3)
add_para('1. 下载 MinIO: https://min.io/download')
add_para('2. 启动：minio.exe server .\\minio-data')
add_para('3. 访问控制台：http://localhost:9001')

doc.add_page_break()

# ============ 第三章 项目结构详解 ============
add_title('第三章 项目结构详解', level=1)

add_title('3.1 整体目录结构', level=2)
add_para('项目采用前后端分离架构，主要目录结构如下：')
add_table(
    ['目录', '说明'],
    [
        ['backend/', '后端 Python 项目'],
        ['backend/app/api/', 'API 路由模块'],
        ['backend/app/core/', '核心配置模块'],
        ['backend/app/models/', '数据库模型'],
        ['backend/app/schemas/', 'Pydantic 模型'],
        ['backend/app/services/', '业务逻辑'],
        ['backend/app/tasks/', 'Celery 任务'],
        ['frontend/', '前端 Vue 项目'],
        ['frontend/src/views/', '页面组件'],
        ['frontend/src/stores/', '状态管理'],
        ['frontend/src/router/', '路由配置'],
        ['docs/', '文档目录'],
    ]
)

add_title('3.2 后端目录详解', level=2)

add_title('3.2.1 API 路由模块', level=3)
add_table(
    ['文件', '说明', '主要接口'],
    [
        ['auth.py', '认证接口', 'login, logout, me'],
        ['users.py', '用户管理', 'CRUD, reset-password'],
        ['projects.py', '项目管理', 'CRUD, members'],
        ['envs.py', '环境管理', 'CRUD, adjust, train'],
        ['models.py', '模型管理', 'upload, download'],
        ['optimization.py', '优化接口', 'evaluate, tasks'],
    ]
)

add_title('3.2.2 核心配置模块', level=3)
add_table(
    ['文件', '说明'],
    [
        ['config.py', '应用配置管理'],
        ['database.py', '数据库连接和会话'],
        ['security.py', '认证授权（JWT、密码）'],
    ]
)

add_title('3.2.3 数据库模型', level=3)
add_para('系统共包含 18 张数据表，核心表包括：')
add_table(
    ['表名', '说明', '主要字段'],
    [
        ['users', '用户表', 'id, username, password_hash, global_role'],
        ['projects', '项目表', 'id, name, description, created_by'],
        ['envs', '环境表', 'id, project_id, name, config, status'],
        ['env_snapshots', '环境快照表', 'id, env_id, config, trigger_type'],
        ['models', '模型表', 'id, project_id, name, type, status'],
        ['optimization_tasks', '优化任务表', 'id, project_id, param_space, status'],
    ]
)

add_title('3.2.4 业务逻辑模块', level=3)
add_table(
    ['文件', '说明'],
    [
        ['env_generator.py', '环境生成服务'],
        ['jsbsim_engine.py', 'JSBSim 飞行模拟引擎封装'],
        ['strategy_engine.py', '策略引擎（规则驱动）'],
        ['evaluator.py', '环境质量评估器'],
        ['optimizer.py', '贝叶斯优化器'],
        ['training_service.py', '训练服务管理'],
        ['ws_manager.py', 'WebSocket 连接管理'],
    ]
)

add_title('3.3 前端目录详解', level=2)

add_title('3.3.1 页面组件', level=3)
add_table(
    ['文件', '说明', '主要功能'],
    [
        ['Login.vue', '登录页', '用户登录表单'],
        ['Envs.vue', '环境管理', '环境配置、3D 预览'],
        ['Monitor.vue', '训练监控', '实时训练曲线'],
        ['Optimization.vue', '优化中心', '环境评估、智能优化'],
        ['Models.vue', '模型库', '模型上传、版本管理'],
        ['Settings.vue', '设置', '用户/成员管理'],
    ]
)

add_title('3.3.2 状态管理', level=3)
add_table(
    ['文件', '说明'],
    [
        ['auth.ts', '认证状态管理（token、用户信息）'],
        ['project.ts', '项目状态管理（项目列表、当前项目）'],
    ]
)

add_title('3.3.3 路由配置', level=3)
add_table(
    ['路径', '页面', '说明'],
    [
        ['/login', 'Login.vue', '登录页'],
        ['/envs', 'Envs.vue', '环境管理'],
        ['/monitor', 'Monitor.vue', '训练监控'],
        ['/optimization', 'Optimization.vue', '优化中心'],
        ['/models', 'Models.vue', '模型库'],
        ['/settings', 'Settings.vue', '设置'],
    ]
)

doc.add_page_break()

# ============ 第四章 开发规范 ============
add_title('第四章 开发规范', level=1)

add_title('4.1 Python 代码规范', level=2)

add_title('4.1.1 代码风格', level=3)
add_para('遵循 PEP 8 规范，使用 4 个空格缩进，行长度限制 88 个字符。')

add_title('4.1.2 命名规范', level=3)
add_table(
    ['类型', '规范', '示例'],
    [
        ['文件名', 'snake_case', 'env_generator.py'],
        ['类名', 'PascalCase', 'EnvGenerator'],
        ['函数名', 'snake_case', 'generate_env'],
        ['变量名', 'snake_case', 'env_config'],
        ['常量名', 'UPPER_SNAKE_CASE', 'MAX_RETRIES'],
    ]
)

add_title('4.1.3 函数文档', level=3)
add_para('每个函数应包含 docstring，说明功能、参数、返回值：')
add_para('def get_user(user_id: str) -> dict: """获取用户信息"""')

add_title('4.2 TypeScript/Vue 代码规范', level=2)

add_title('4.2.1 Vue 组件规范', level=3)
add_para('使用 <script setup> 语法，使用 Composition API，组件名使用 PascalCase。')

add_title('4.2.2 TypeScript 规范', level=3)
add_para('使用 interface 定义对象类型，使用 type 定义联合类型，函数声明返回类型。')

add_title('4.3 Git 规范', level=2)

add_title('4.3.1 分支命名', level=3)
add_table(
    ['分支', '说明'],
    [
        ['main', '主分支'],
        ['develop', '开发分支'],
        ['feature/xxx', '功能分支'],
        ['fix/xxx', '修复分支'],
        ['release/x.x.x', '发布分支'],
    ]
)

add_title('4.3.2 提交信息格式', level=3)
add_para('<type>(<scope>): <subject>')
add_para('Type 类型：feat(新功能)、fix(修复)、docs(文档)、refactor(重构)、test(测试)')

doc.add_page_break()

# ============ 第五章 调试与测试 ============
add_title('第五章 调试与测试', level=1)

add_title('5.1 后端调试', level=2)

add_title('5.1.1 启动服务', level=3)
add_para('cd backend && .\\venv\\Scripts\\activate && python run.py')
add_para('或使用 uvicorn：uvicorn app.main:app --reload --port 8000')

add_title('5.1.2 API 测试', level=3)
add_para('启动后访问 API 文档：http://localhost:8000/docs')
add_para('使用 curl 或 httpie 测试接口。')

add_title('5.1.3 日志查看', level=3)
add_para('后端日志输出到控制台，可查看请求和错误信息。')

add_title('5.2 前端调试', level=2)

add_title('5.2.1 启动开发服务器', level=3)
add_para('cd frontend && npm run dev')
add_para('访问 http://localhost:5173')

add_title('5.2.2 浏览器开发者工具', level=3)
add_para('按 F12 打开开发者工具，使用 Console、Network、Elements 面板调试。')

add_title('5.2.3 Vue DevTools', level=3)
add_para('安装 Vue DevTools 浏览器扩展，查看组件树和状态。')

add_title('5.3 数据库调试', level=2)

add_title('5.3.1 使用 psql', level=3)
add_para('psql -U postgres -d fltect')
add_para('\\dt  查看所有表')
add_para('\\d users  查看表结构')

add_title('5.3.2 图形化工具', level=3)
add_para('推荐使用 pgAdmin、DBeaver 或 Navicat 进行数据库管理。')

doc.add_page_break()

# ============ 第六章 部署指南 ============
add_title('第六章 部署指南', level=1)

add_title('6.1 Docker 部署', level=2)

add_title('6.1.1 前提条件', level=3)
add_para('安装 Docker Desktop: https://www.docker.com/products/docker-desktop/')

add_title('6.1.2 启动服务', level=3)
add_para('docker compose up -d --build')

add_title('6.1.3 服务端口', level=3)
add_table(
    ['服务', '端口', '说明'],
    [
        ['frontend', '80', '前端页面'],
        ['backend', '8000', '后端 API'],
        ['postgres', '5432', '数据库'],
        ['redis', '6379', '缓存'],
        ['minio', '9000/9001', '对象存储'],
    ]
)

add_title('6.2 本地部署', level=2)

add_title('6.2.1 启动后端', level=3)
add_para('cd backend && .\\venv\\Scripts\\activate && python run.py')

add_title('6.2.2 启动前端', level=3)
add_para('cd frontend && npm run dev')

add_title('6.2.3 一键启动', level=3)
add_para('Windows: .\\start_dev.ps1 或 start.bat')

doc.add_page_break()

# ============ 第七章 常见问题 ============
add_title('第七章 常见问题', level=1)

add_title('7.1 依赖安装问题', level=2)
add_para('Q: pip install 失败', bold=True)
add_para('A: 使用国内镜像 pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple')
add_para('Q: npm install 失败', bold=True)
add_para('A: 清除缓存 npm cache clean --force，使用淘宝镜像')

add_title('7.2 数据库问题', level=2)
add_para('Q: 连接失败', bold=True)
add_para('A: 检查 PostgreSQL 服务是否运行，检查端口和密码配置。')
add_para('Q: 表不存在', bold=True)
add_para('A: 启动应用会自动创建表，或使用 alembic upgrade head')

add_title('7.3 端口问题', level=2)
add_para('Q: 端口被占用', bold=True)
add_para('A: 使用 netstat -ano | findstr :8000 查找进程，taskkill /PID 结束进程。')

add_title('7.4 认证问题', level=2)
add_para('Q: JWT Token 无效', bold=True)
add_para('A: 检查 Token 是否过期，检查 JWT_SECRET_KEY 配置，检查请求头格式。')

# ============ 附录 ============
doc.add_page_break()
add_title('附录', level=1)

add_title('A. 默认账号', level=2)
add_table(
    ['用户名', '密码', '角色'],
    [
        ['admin', 'admin123', '管理员'],
    ]
)

add_title('B. 环境变量配置', level=2)
add_para('DATABASE_URL=postgresql+asyncpg://postgres:password@localhost:5432/fltect')
add_para('REDIS_URL=redis://localhost:6379/0')
add_para('JWT_SECRET_KEY=your-secret-key')

add_title('C. 常用命令', level=2)
add_table(
    ['命令', '说明'],
    [
        ['python run.py', '启动后端'],
        ['npm run dev', '启动前端'],
        ['docker compose up -d', 'Docker 启动'],
        ['docker compose down', 'Docker 停止'],
        ['pip install -r requirements.txt', '安装依赖'],
        ['npm install', '安装前端依赖'],
    ]
)

# 保存文档
output_path = r'c:\Users\mujunze\Desktop\Flight-Test-Environment-Construction-System-0f10903b41da7005c094e6abd7b1f142f9c58bda\docs\开发手册.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
